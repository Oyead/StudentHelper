import os
from abc import ABC, abstractmethod
from PyQt5.QtWidgets import QFileDialog


# ------------------- File Selection -------------------
def select_file():
    path, _ = QFileDialog.getOpenFileName(
        None,
        "Select a file",
        "",
        "All Files (*);;Word Files (*.docx);;PDF Files (*.pdf);;Text Files (*.txt);;HTML Files (*.html *.htm)"
    )
    return path if path else None


# ------------------- Base Converter -------------------
class ConverterBase(ABC):
    @abstractmethod
    def convert(self, input_path, output_dir):
        pass


# ------------------- Word Converters -------------------
class WordToPDFConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from docx2pdf import convert as docx2pdf_convert
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
        docx2pdf_convert(input_path, out_path)


class WordToTXTConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from docx import Document
        doc = Document(input_path)
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".txt")
        with open(out_path, "w", encoding="utf-8") as f:
            for p in doc.paragraphs:
                f.write(p.text + "\n")


class WordToHTMLConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        import mammoth
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".html")
        with open(input_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(html)


# ------------------- PDF Converters -------------------
class PDFToTXTConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        import PyPDF2
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".txt")
        text = ""
        with open(input_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text)


class PDFToHTMLConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from pdf2docx import Converter
        import mammoth

        # Step 1: Convert PDF → Word
        word_path = os.path.join(
            output_dir, os.path.splitext(os.path.basename(input_path))[0] + "_temp.docx"
        )
        cv = Converter(input_path)
        cv.convert(word_path, start=0, end=None)
        cv.close()

        # Step 2: Convert Word → HTML
        with open(word_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value  # HTML content
            # images are converted as base64 automatically
            # if you want them saved as separate files, mammoth supports that

        out_path = os.path.join(
            output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".html"
        )
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        # Optional: remove temporary Word file
        os.remove(word_path)
        print(f"Saved HTML: {out_path}")


class PDFToWordConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from pdf2docx import Converter

        out_path = os.path.join(
            output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".docx"
        )
        cv = Converter(input_path)
        cv.convert(out_path, start=0, end=None)
        cv.close()
        print(f"Saved Word: {out_path}")
        return out_path


# ------------------- TXT Converters -------------------
class TXTToPDFConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                pdf.multi_cell(0, 8, line.rstrip())
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
        pdf.output(out_path)


class TXTToWordConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from docx import Document
        doc = Document()
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                doc.add_paragraph(line.rstrip())
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".docx")
        doc.save(out_path)


class TXTToHTMLConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".html")
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("<html><body><pre>\n")
            for line in lines:
                f.write(line)
            f.write("</pre></body></html>")


# ------------------- HTML Converters -------------------
class HTMLToTXTConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from bs4 import BeautifulSoup
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".txt")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text)


class HTMLToWordConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        from bs4 import BeautifulSoup
        from docx import Document
        doc = Document()
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f, "html.parser")
        for element in soup.body.descendants:
            if element.name == "h1":
                doc.add_heading(element.get_text(), level=1)
            elif element.name == "h2":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name == "li":
                doc.add_paragraph(element.get_text(), style="List Bullet")
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".docx")
        doc.save(out_path)


class HTMLToPDFConverter(ConverterBase):
    def convert(self, input_path, output_dir):
        import pdfkit
        out_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
        pdfkit.from_file(input_path, out_path)


# ------------------- Conversion Manager -------------------
class ConversionManager:
    converters = {
        ".docx": {"pdf": WordToPDFConverter(), "txt": WordToTXTConverter(), "html": WordToHTMLConverter()},
        ".pdf": {"txt": PDFToTXTConverter(), "html": PDFToHTMLConverter(), "docx": PDFToWordConverter()},
        ".txt": {"pdf": TXTToPDFConverter(), "html": TXTToHTMLConverter(), "docx": TXTToWordConverter()},
        ".html": {"pdf": HTMLToPDFConverter(), "txt": HTMLToTXTConverter(), "docx": HTMLToWordConverter()},
    }

    @classmethod
    def convert(cls, input_path, output_dir, target_format):
        ext = os.path.splitext(input_path)[1].lower()
        if ext not in cls.converters:
            raise ValueError(f"Unsupported input file: {ext}")
        if target_format not in cls.converters[ext]:
            raise ValueError(f"Unsupported output format: {target_format}")
        cls.converters[ext][target_format].convert(input_path, output_dir)


# ------------------- Public API -------------------
def convert(input_path, output_dir, target_format):
    ConversionManager.convert(input_path, output_dir, target_format)
